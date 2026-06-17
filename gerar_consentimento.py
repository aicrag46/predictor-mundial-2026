"""
Gera o documento Word de consentimento para o Predictor Mundial 2026.
UMA FOLHA por jogador — tudo compactado para caber em A4 Portrait.
"""
import json, datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Dados ─────────────────────────────────────────────────────────────────────
with open("/tmp/predictor_data.json", encoding="utf-8") as f:
    dados = json.load(f)

participantes = dados["participantes"]
jogos         = dados["jogos"]
prognosticos  = dados["prognosticos"]
grupos_order  = list(dict.fromkeys(j["grupo"] for j in jogos))
hoje          = datetime.date.today().strftime("%d/%m/%Y")

# ── Cores ─────────────────────────────────────────────────────────────────────
AZUL   = RGBColor(0x1a, 0x3a, 0x6e)
AZUL_L = RGBColor(0xd6, 0xe4, 0xf0)
AMAR   = RGBColor(0xf5, 0xa6, 0x23)
VERM   = RGBColor(0xcc, 0x22, 0x22)
BRNC   = RGBColor(0xff, 0xff, 0xff)
CZ1    = RGBColor(0xf2, 0xf5, 0xf9)
CZ2    = RGBColor(0x55, 0x55, 0x55)

def hx(rgb): return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

# ── Helpers XML ───────────────────────────────────────────────────────────────
def cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hx(rgb))
    tcPr.append(shd)

def cell_no_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    b    = OxmlElement("w:tcBorders")
    for s in ("top","bottom","left","right","insideH","insideV"):
        el = OxmlElement(f"w:{s}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "FFFFFF")
        b.append(el)
    tcPr.append(b)

def cell_margins(cell, top=0, bottom=0, left=36, right=36):
    """Margens internas da célula em twips (1cm=567)."""
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def para_shd(para, rgb):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hx(rgb))
    pPr.append(shd)

def set_line_spacing(para, lines=1.0):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = Pt(lines * 12)

def p(para, text, bold=False, italic=False, sz=8, color=None):
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size      = Pt(sz)
    run.font.color.rgb = color or RGBColor(0x11,0x11,0x11)
    return run

def fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
    pf = para.paragraph_format
    pf.alignment    = align
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = Pt(11)

def merge_row(table, row_idx, text, bg, sz=7, color=BRNC, bold=True):
    """Funde todas as células de uma linha e escreve texto."""
    row   = table.rows[row_idx]
    merged = row.cells[0]
    for i in range(1, len(row.cells)):
        merged = merged.merge(row.cells[i])
    cell_bg(merged, bg)
    cell_margins(merged, top=28, bottom=28, left=60, right=0)
    pp = merged.paragraphs[0]
    fmt(pp)
    p(pp, text, bold=bold, sz=sz, color=color)
    return merged

# ── Documento ─────────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(8)

for sec in doc.sections:
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.0)
    sec.right_margin  = Cm(1.0)
    sec.top_margin    = Cm(0.9)
    sec.bottom_margin = Cm(0.9)
    # Sem cabeçalho/rodapé de secção
    sec.header_distance = Cm(0.5)
    sec.footer_distance = Cm(0.5)

# Área útil: 21 - 2 = 19cm de largura
W = 19.0  # cm

# ── Uma página por jogador ────────────────────────────────────────────────────
for idx_p, nome in enumerate(participantes):
    preds = prognosticos.get(nome, {})

    # ─────────────────────────────────────────────────────
    # 1. TÍTULO + BANNER (tabela 1×2)
    # ─────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=2)
    hdr.style    = "Table Grid"
    hdr.autofit  = False
    hdr.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Linha 0: título do documento (span total)
    r0 = hdr.rows[0]
    title_cell = r0.cells[0].merge(r0.cells[1])
    cell_bg(title_cell, AZUL)
    cell_margins(title_cell, top=60, bottom=40, left=120, right=60)
    tp = title_cell.paragraphs[0]
    fmt(tp); p(tp, "🏆  Predictor Parque Biológico — Mundial 2026   ",
               bold=True, sz=11, color=BRNC)
    p(tp, "· Folha de Verificação de Prognósticos", sz=8, color=RGBColor(0xcc,0xdd,0xff))

    # Linha 1: nome do jogador | data
    r1 = hdr.rows[1]
    r1.cells[0].width = Cm(W * 0.68)
    r1.cells[1].width = Cm(W * 0.32)
    cell_bg(r1.cells[0], RGBColor(0x0f, 0x28, 0x55))
    cell_bg(r1.cells[1], RGBColor(0x0f, 0x28, 0x55))
    cell_margins(r1.cells[0], top=50, bottom=50, left=120, right=60)
    cell_margins(r1.cells[1], top=50, bottom=50, left=60, right=120)

    np = r1.cells[0].paragraphs[0]; fmt(np)
    p(np, "Jogador:  ", sz=8, color=RGBColor(0xaa,0xcc,0xff))
    p(np, nome, bold=True, sz=12, color=AMAR)

    dp = r1.cells[1].paragraphs[0]; fmt(dp, align=WD_ALIGN_PARAGRAPH.RIGHT)
    p(dp, f"Data: {hoje}", sz=8, color=RGBColor(0xaa,0xcc,0xff))

    # ─────────────────────────────────────────────────────
    # 2. INSTRUÇÃO (1 linha)
    # ─────────────────────────────────────────────────────
    inst = doc.add_paragraph()
    fmt(inst, sb=3, sa=2)
    para_shd(inst, RGBColor(0xfe, 0xf9, 0xe7))
    p(inst, "⚠ ", sz=7.5, color=RGBColor(0xb7, 0x7d, 0x00))
    p(inst, "Instruções: ", bold=True, sz=7.5, color=AZUL)
    p(inst, "Compare os prognósticos abaixo com a fotocópia da folha original. "
            "Registe na secção de observações qualquer resultado que não coincida. "
            "Assine no final para confirmar.", sz=7.5, color=CZ2)

    # ─────────────────────────────────────────────────────
    # 3. TÍTULO FASE DE GRUPOS
    # ─────────────────────────────────────────────────────
    gs_hdr = doc.add_paragraph()
    fmt(gs_hdr, sb=3, sa=1)
    p(gs_hdr, "⚽  FASE DE GRUPOS — Prognósticos registados na Base de Dados",
      bold=True, sz=8, color=AZUL)

    # ─────────────────────────────────────────────────────
    # 4. TABELA DE GRUPOS (2 colunas × 6 grupos)
    # ─────────────────────────────────────────────────────
    # Tabela exterior sem bordas
    outer = doc.add_table(rows=1, cols=2)
    outer.style   = "Table Grid"
    outer.autofit = False
    outer.alignment = WD_TABLE_ALIGNMENT.LEFT

    COL_W = W / 2  # 9.5cm por coluna
    outer.columns[0].width = Cm(COL_W)
    outer.columns[1].width = Cm(COL_W)
    for oc in outer.rows[0].cells:
        cell_no_border(oc)
        cell_margins(oc, top=0, bottom=0, left=0, right=20)
        oc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Larguras colunas internas: Cód | Casa | Fora | Prev.
    GCW = [0.55, 3.65, 3.65, 0.95]  # total ≈ 8.8 (< 9.5)

    metades = [grupos_order[:6], grupos_order[6:]]

    for ci, grp_list in enumerate(metades):
        ocell = outer.rows[0].cells[ci]

        for gi, grupo in enumerate(grp_list):
            jgs = [j for j in jogos if j["grupo"] == grupo]
            n_rows = 1 + 1 + len(jgs)  # título + cabeçalho + jogos

            gt = ocell.add_table(rows=n_rows, cols=4)
            gt.style   = "Table Grid"
            gt.autofit = False
            for row in gt.rows:
                for c_idx, w in enumerate(GCW):
                    row.cells[c_idx].width = Cm(w)

            # Fila 0: título do grupo (merged)
            title_merged = gt.rows[0].cells[0]
            for x in range(1, 4):
                title_merged = title_merged.merge(gt.rows[0].cells[x])
            cell_bg(title_merged, AZUL)
            cell_margins(title_merged, top=22, bottom=22, left=60, right=0)
            gtp = title_merged.paragraphs[0]; fmt(gtp)
            p(gtp, f"GRUPO {grupo}", bold=True, sz=6.5, color=BRNC)

            # Fila 1: cabeçalhos das colunas
            for x, hd in enumerate(["Cód", "Casa", "Fora", "Prev."]):
                hcell = gt.rows[1].cells[x]
                cell_bg(hcell, AZUL_L)
                cell_margins(hcell, top=18, bottom=18, left=50, right=10)
                hp = hcell.paragraphs[0]; fmt(hp)
                p(hp, hd, bold=True, sz=6, color=AZUL)

            # Filas 2+: jogos
            for ri, jogo in enumerate(jgs):
                pred     = preds.get(jogo["codigo"])
                pred_str = f"{pred['casa']}–{pred['fora']}" if pred else "—"
                row_bg   = CZ1 if ri % 2 == 0 else BRNC

                vals  = [jogo["codigo"], jogo["casa"], jogo["fora"], pred_str]
                bolds = [True, False, False, True]
                cols  = [AZUL, RGBColor(0x22,0x22,0x22), RGBColor(0x22,0x22,0x22), VERM]

                for x, (val, bd, cl) in enumerate(zip(vals, bolds, cols)):
                    gc = gt.rows[ri + 2].cells[x]
                    cell_bg(gc, row_bg)
                    cell_margins(gc, top=16, bottom=16, left=50, right=10)
                    gp2 = gc.paragraphs[0]; fmt(gp2)
                    p(gp2, val, bold=bd, sz=6.5, color=cl)

            # Pequeno espaço após cada grupo (excepto o último)
            if gi < len(grp_list) - 1:
                sp = ocell.add_paragraph()
                fmt(sp, sb=1, sa=0)

    # ─────────────────────────────────────────────────────
    # 5. OBSERVAÇÕES
    # ─────────────────────────────────────────────────────
    obs_hdr = doc.add_paragraph()
    fmt(obs_hdr, sb=4, sa=1)
    p(obs_hdr, "📋  OBSERVAÇÕES — Resultados que não coincidem com a folha original",
      bold=True, sz=8, color=AZUL)

    N_OBS = 5
    obs_t = doc.add_table(rows=N_OBS + 1, cols=4)
    obs_t.style   = "Table Grid"
    obs_t.autofit = False
    obs_t.alignment = WD_TABLE_ALIGNMENT.LEFT

    OW = [2.5, 4.0, 4.0, 8.5]  # total 19cm
    for row in obs_t.rows:
        for x, w in enumerate(OW):
            row.cells[x].width = Cm(w)

    for x, hd in enumerate(["Jogo (código)", "Prognóstico na folha", "Prognóstico na BD", "Nota / Decisão tomada"]):
        hcell = obs_t.rows[0].cells[x]
        cell_bg(hcell, AZUL)
        cell_margins(hcell, top=30, bottom=30, left=60, right=10)
        hp = hcell.paragraphs[0]; fmt(hp)
        p(hp, hd, bold=True, sz=7, color=BRNC)

    for ri in range(1, N_OBS + 1):
        bg = CZ1 if ri % 2 == 0 else BRNC
        for x in range(4):
            gc = obs_t.rows[ri].cells[x]
            cell_bg(gc, bg)
            cell_margins(gc, top=42, bottom=42, left=60, right=10)
            obs_t.rows[ri].cells[x].paragraphs[0].add_run("")

    # ─────────────────────────────────────────────────────
    # 6. DECLARAÇÃO + ASSINATURA
    # ─────────────────────────────────────────────────────
    decl_hdr = doc.add_paragraph()
    fmt(decl_hdr, sb=4, sa=1)
    p(decl_hdr, "📝  DECLARAÇÃO DE CONSENTIMENTO",
      bold=True, sz=8, color=AZUL)

    # Caixa da declaração
    db = doc.add_table(rows=1, cols=1)
    db.style   = "Table Grid"
    db.autofit = False
    db.columns[0].width = Cm(W)
    cell_bg(db.rows[0].cells[0], RGBColor(0xee, 0xf4, 0xfb))
    cell_margins(db.rows[0].cells[0], top=50, bottom=50, left=100, right=100)

    dp = db.rows[0].cells[0].paragraphs[0]; fmt(dp)
    p(dp, f"Eu, ", sz=8)
    p(dp, nome, bold=True, sz=8)
    p(dp, ", declaro que revi os prognósticos acima e confirmo que são os mesmos que preenchi na "
          "folha original, com exceção das discrepâncias registadas nas observações. "
          "Aceito que os valores na Base de Dados sejam utilizados para a classificação final do Predictor.", sz=8)

    # Assinatura + data (tabela sem bordas, 2 colunas)
    sig = doc.add_table(rows=2, cols=2)
    sig.style   = "Table Grid"
    sig.autofit = False
    sig.columns[0].width = Cm(W * 0.62)
    sig.columns[1].width = Cm(W * 0.38)
    for ri in range(2):
        for ci in range(2):
            cell_no_border(sig.rows[ri].cells[ci])

    # Linha de assinatura
    sp0 = sig.rows[0].cells[0].paragraphs[0]
    fmt(sp0, sb=18, sa=0)
    p(sp0, "_" * 58, sz=8.5)

    dp0 = sig.rows[0].cells[1].paragraphs[0]
    fmt(dp0, sb=18, sa=0)
    p(dp0, "___ / ___ / 2026", sz=8.5)

    # Labels
    sp1 = sig.rows[1].cells[0].paragraphs[0]
    fmt(sp1, sb=1, sa=0)
    p(sp1, f"Assinatura de {nome}", italic=True, sz=7.5, color=CZ2)

    dp1 = sig.rows[1].cells[1].paragraphs[0]
    fmt(dp1, sb=1, sa=0)
    p(dp1, "Data", italic=True, sz=7.5, color=CZ2)

    # ─────────────────────────────────────────────────────
    # 7. RODAPÉ
    # ─────────────────────────────────────────────────────
    foot = doc.add_paragraph()
    fmt(foot, align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=0)
    p(foot, f"Predictor Parque Biológico — Mundial 2026  ·  {nome}  ·  {hoje}  ·  "
            f"Pág. {idx_p+1}/{len(participantes)}",
      sz=6.5, color=RGBColor(0xaa,0xaa,0xaa))

    # Quebra de página (excepto no último jogador)
    if idx_p < len(participantes) - 1:
        br = doc.add_paragraph()
        fmt(br)
        br.add_run().add_break(
            __import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE
        )

# ── Guardar ───────────────────────────────────────────────────────────────────
OUT = "/Users/joaogarcia/Projects/predictor-mundial-2026/Consentimento_Predictor_Mundial_2026.docx"
doc.save(OUT)
print(f"✅  {OUT}")
print(f"    {len(participantes)} páginas · uma por jogador")
