"""
Gera o Regulamento Oficial do Predictor Parque Biológico — Mundial 2026.
Documento Word completo, detalhado, com exemplos, penalizações e assinatura.
"""
import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Cores ─────────────────────────────────────────────────────────────────────
AZUL   = RGBColor(0x1a, 0x3a, 0x6e)
AZUL_L = RGBColor(0xd6, 0xe4, 0xf0)
AMAR   = RGBColor(0xb8, 0x6e, 0x00)
VERM   = RGBColor(0xcc, 0x22, 0x22)
VERDE  = RGBColor(0x1a, 0x6e, 0x3a)
BRNC   = RGBColor(0xff, 0xff, 0xff)
CZ1    = RGBColor(0xf2, 0xf5, 0xf9)
CZ2    = RGBColor(0x44, 0x44, 0x44)
PRETO  = RGBColor(0x11, 0x11, 0x11)

hoje = datetime.date.today().strftime("%d de %B de %Y")

def hx(rgb): return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"

# ── Helpers ───────────────────────────────────────────────────────────────────
def cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hx(rgb)); tcPr.append(shd)

def cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for s, v in (("top",top),("bottom",bottom),("left",left),("right",right)):
        el = OxmlElement(f"w:{s}"); el.set(qn("w:w"), str(v)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)

def fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4, indent=0):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = Pt(13.5)
    if indent: pf.left_indent = Cm(indent)

def run(para, text, bold=False, italic=False, sz=10, color=None):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(sz)
    r.font.color.rgb = color or PRETO
    return r

# ── Documento ─────────────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

for sec in doc.sections:
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

# ──────────────────────────────────────────────────────────────────────────────
# CAPA
# ──────────────────────────────────────────────────────────────────────────────
for _ in range(4): doc.add_paragraph()

capa_trophy = doc.add_paragraph()
fmt(capa_trophy, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
run(capa_trophy, "🏆", sz=36)

capa1 = doc.add_paragraph()
fmt(capa1, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(capa1, "PREDICTOR PARQUE BIOLÓGICO", bold=True, sz=20, color=AZUL)

capa2 = doc.add_paragraph()
fmt(capa2, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(capa2, "MUNDIAL 2026", bold=True, sz=26, color=AZUL)

capa3 = doc.add_paragraph()
fmt(capa3, align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)
run(capa3, "⚽  REGULAMENTO OFICIAL  ⚽", bold=True, sz=14, color=AMAR)

# Linha decorativa
sep = doc.add_paragraph()
fmt(sep, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)
run(sep, "━" * 45, sz=9, color=AZUL)

capa4 = doc.add_paragraph()
fmt(capa4, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(capa4, "10 participantes  ·  72 jogos  ·  720 prognósticos", sz=11, color=CZ2)

capa5 = doc.add_paragraph()
fmt(capa5, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(capa5, "Fase de Grupos + Mata-Mata Completo", sz=11, color=CZ2)

for _ in range(6): doc.add_paragraph()

capa_ver = doc.add_paragraph()
fmt(capa_ver, align=WD_ALIGN_PARAGRAPH.CENTER, sa=1)
run(capa_ver, "Versão 1.0 — Edição FIFA World Cup 2026", italic=True, sz=9, color=CZ2)

capa_data = doc.add_paragraph()
fmt(capa_data, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(capa_data, f"Emitido em {hoje}", italic=True, sz=9, color=CZ2)

# Quebra para página 2
br = doc.add_paragraph(); fmt(br)
br.add_run().add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

# ──────────────────────────────────────────────────────────────────────────────
# ÍNDICE
# ──────────────────────────────────────────────────────────────────────────────
idx_t = doc.add_paragraph()
fmt(idx_t, sa=6)
run(idx_t, "ÍNDICE", bold=True, sz=14, color=AZUL)

artigos = [
    ("Preâmbulo", ""),
    ("Artigo 1.º",  "Objeto e Âmbito"),
    ("Artigo 2.º",  "Participantes"),
    ("Artigo 3.º",  "Inscrição e Entrega dos Prognósticos"),
    ("Artigo 4.º",  "Fases do Torneio"),
    ("Artigo 5.º",  "Pontuação — Fase de Grupos"),
    ("Artigo 6.º",  "Pontuação — Mata-Mata"),
    ("Artigo 7.º",  "Tabela de Pontuação Progressiva"),
    ("Artigo 8.º",  "Classificação Geral"),
    ("Artigo 9.º",  "Critérios de Desempate"),
    ("Artigo 10.º", "Prémio e Penalização — O Jantar"),
    ("Artigo 11.º", "Obrigações dos Últimos 5 Classificados"),
    ("Artigo 12.º", "Ausência ao Jantar"),
    ("Artigo 13.º", "Suspense — Revelação dos Resultados no Jantar"),
    ("Artigo 14.º", "Reclamações e Resolução de Disputas"),
    ("Artigo 15.º", "Base de Dados e Verificação dos Prognósticos"),
    ("Artigo 16.º", "Alteração de Prognósticos"),
    ("Artigo 17.º", "Jogos Cancelados, Adiados ou com Resultado Anulado"),
    ("Artigo 18.º", "Conduta e Fair-Play"),
    ("Artigo 19.º", "Privacidade e Uso dos Dados"),
    ("Artigo 20.º", "Disposições Finais"),
]
for art, titulo in artigos:
    ip = doc.add_paragraph()
    fmt(ip, sa=1, indent=0)
    run(ip, f"{art}  ", bold=True, sz=9.5, color=AZUL)
    if titulo:
        run(ip, f"— {titulo}", sz=9.5, color=CZ2)

br2 = doc.add_paragraph(); fmt(br2)
br2.add_run().add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

# ──────────────────────────────────────────────────────────────────────────────
# HELPERS de formatação do corpo
# ──────────────────────────────────────────────────────────────────────────────
def section_title(doc, text, sb=12):
    p2 = doc.add_paragraph()
    fmt(p2, sb=sb, sa=3)
    run(p2, text, bold=True, sz=13, color=AZUL)
    # Linha sob o título
    line = doc.add_paragraph()
    fmt(line, sb=0, sa=6)
    run(line, "─" * 80, sz=5, color=AZUL)
    return p2

def art_title(doc, artigo, titulo, sb=10):
    p2 = doc.add_paragraph()
    fmt(p2, sb=sb, sa=2)
    run(p2, f"{artigo} — {titulo}", bold=True, sz=11, color=AZUL)
    return p2

def body(doc, text, indent=0, sb=0, sa=3):
    p2 = doc.add_paragraph()
    fmt(p2, sa=sa, sb=sb, indent=indent)
    run(p2, text, sz=10)
    return p2

def bullet(doc, text, indent=0.7, symbol="•"):
    p2 = doc.add_paragraph()
    fmt(p2, sa=1, indent=indent)
    run(p2, f"{symbol}  ", bold=True, sz=10, color=AZUL)
    run(p2, text, sz=10)
    return p2

def numbered(doc, num, text, indent=0.7):
    p2 = doc.add_paragraph()
    fmt(p2, sa=2, indent=indent)
    run(p2, f"{num}.  ", bold=True, sz=10, color=AZUL)
    run(p2, text, sz=10)
    return p2

def nota(doc, text, sb=4, sa=4):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"; tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    c = tbl.rows[0].cells[0]
    cell_bg(c, RGBColor(0xfe, 0xf9, 0xe7))
    cell_margins(c, top=80, bottom=80, left=100, right=100)
    np = c.paragraphs[0]; fmt(np, sb=0, sa=0)
    run(np, "⚠  Nota:  ", bold=True, sz=9.5, color=AMAR)
    run(np, text, sz=9.5, color=CZ2)
    sp = doc.add_paragraph(); fmt(sp, sb=sb, sa=sa)
    return tbl

def exemplo(doc, titulo, linhas, sb=4):
    tbl = doc.add_table(rows=len(linhas)+1, cols=1)
    tbl.style = "Table Grid"; tbl.autofit = False
    tbl.columns[0].width = Cm(16)
    # Cabeçalho
    hc = tbl.rows[0].cells[0]; cell_bg(hc, AZUL)
    cell_margins(hc, top=50, bottom=50, left=100, right=100)
    hp = hc.paragraphs[0]; fmt(hp, sa=0)
    run(hp, f"📌  Exemplo: {titulo}", bold=True, sz=9.5, color=BRNC)
    # Linhas
    for ri, linha in enumerate(linhas):
        bc = tbl.rows[ri+1].cells[0]
        cell_bg(bc, CZ1 if ri%2==0 else BRNC)
        cell_margins(bc, top=40, bottom=40, left=120, right=100)
        lp = bc.paragraphs[0]; fmt(lp, sa=0)
        if isinstance(linha, tuple):
            run(lp, linha[0], bold=True, sz=9.5, color=AZUL)
            run(lp, "  " + linha[1], sz=9.5)
        else:
            run(lp, linha, sz=9.5)
    sp = doc.add_paragraph(); fmt(sp, sb=sb, sa=0)
    return tbl

def tabela_pontuacao(doc):
    dados_tbl = [
        ("Fase de Grupos",  "5",  "2",  "1",  "—",  "5"),
        ("R32 — 16 avos",   "7",  "3",  "1",  "3",  "10"),
        ("R16 — Oitavos",   "10", "4",  "2",  "5",  "15"),
        ("QF — Quartos",    "15", "6",  "3",  "10", "25"),
        ("SF — Meias",      "20", "8",  "4",  "15", "35"),
        ("3.º/4.º Lugar",   "25", "10", "5",  "15", "40"),
        ("🏆 Final",        "35", "15", "6",  "15", "50"),
    ]
    headers = ["Fase", "✅ Exato", "🔵 VE", "🟡 Golos", "🟢 Apurado", "🏅 Máx."]
    tbl = doc.add_table(rows=len(dados_tbl)+1, cols=6)
    tbl.style = "Table Grid"; tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [4.5, 2.0, 2.0, 2.0, 2.5, 2.0]
    for row in tbl.rows:
        for ci, w in enumerate(widths): row.cells[ci].width = Cm(w)
    for ci, h in enumerate(headers):
        hc = tbl.rows[0].cells[ci]; cell_bg(hc, AZUL)
        cell_margins(hc, top=50, bottom=50)
        hp = hc.paragraphs[0]; fmt(hp, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
        run(hp, h, bold=True, sz=8.5, color=BRNC)
    for ri, row_data in enumerate(dados_tbl):
        bg = RGBColor(0xf5, 0xa6, 0x23) if "Final" in row_data[0] else (CZ1 if ri%2==0 else BRNC)
        for ci, val in enumerate(row_data):
            rc = tbl.rows[ri+1].cells[ci]
            cell_bg(rc, bg)
            cell_margins(rc, top=45, bottom=45)
            rp = rc.paragraphs[0]
            al = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            fmt(rp, align=al, sa=0)
            color = AZUL if ci == 0 else (VERDE if ci == 5 else PRETO)
            bold  = ci == 0 or ci == 5
            run(rp, val, bold=bold, sz=9, color=color)
    sp = doc.add_paragraph(); fmt(sp, sb=4, sa=4)


# ──────────────────────────────────────────────────────────────────────────────
# PREÂMBULO
# ──────────────────────────────────────────────────────────────────────────────
section_title(doc, "Preâmbulo")
body(doc, "O Predictor Parque Biológico — Mundial 2026 é um torneio de prognósticos desportivos "
     "organizado de forma informal entre amigos e colegas, com o propósito de tornar o "
     "acompanhamento do Campeonato do Mundo FIFA 2026 mais emocionante, competitivo e divertido.")
body(doc, "O presente Regulamento estabelece todas as regras, critérios de pontuação, procedimentos "
     "de reclamação, penalizações e disposições gerais que regem a competição. "
     "A participação implica a aceitação integral e incondicional de todas as normas aqui descritas.")
body(doc, "Em caso de dúvida sobre a interpretação de qualquer artigo, a decisão final cabe ao "
     "organizador, João Garcia, cujo julgamento é soberano e irrecorrível — salvo disposição "
     "expressa em contrário neste regulamento.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 1.º", "Objeto e Âmbito")
body(doc, "1.1  O presente regulamento rege o torneio de prognósticos denominado "
     "Predictor Parque Biológico — Mundial 2026, doravante designado apenas por «Predictor».")
body(doc, "1.2  O Predictor abrange todos os jogos oficiais do Campeonato do Mundo FIFA 2026, "
     "incluindo a fase de grupos (48 jogos), o play-off intercontinental (16 jogos de eliminação "
     "direta) e a fase a eliminar (R32 até à Final), num total de 72 jogos.")
body(doc, "1.3  O Predictor é de participação exclusivamente voluntária, sem qualquer taxa de "
     "inscrição em dinheiro, sem prémios monetários e sem fins lucrativos.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 2.º", "Participantes")
body(doc, "2.1  O Predictor conta com 10 participantes inscritos, listados por ordem alfabética:")
participantes = ["André", "Hélio", "João Garcia", "José Silva", "Luis Félix",
                 "Miguel Fernandes", "Paulo Afonso", "Ricardo Diegues",
                 "Rúben Gonçalves", "Vítor Fernandes"]
for i, nome in enumerate(participantes, 1):
    bullet(doc, nome)
body(doc, "2.2  A lista de participantes é definitiva e encerrada. Não são aceites inscrições "
     "adicionais após o arranque do torneio.", sb=4)
body(doc, "2.3  A participação é pessoal e intransferível. Nenhum participante pode ceder ou "
     "alterar a sua posição no torneio a favor de terceiros.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 3.º", "Inscrição e Entrega dos Prognósticos")
body(doc, "3.1  Cada participante preencheu previamente uma folha física com os seus prognósticos "
     "para todos os 72 jogos da fase de grupos e, na medida do possível, os prognósticos "
     "para a fase a eliminar (mata-mata).")
body(doc, "3.2  Os prognósticos da fase de grupos foram registados na Base de Dados (BD) digital "
     "do Predictor antes do início do Campeonato do Mundo. Esta BD é a fonte oficial e definitiva "
     "de toda a pontuação.")
body(doc, "3.3  Para os jogos do mata-mata, os prognósticos são recolhidos digitalmente. "
     "O organizador envia a cada participante, por mensagem digital (WhatsApp ou equivalente), "
     "os jogos de cada fase eliminatória antes do seu início. Cada participante responde "
     "com os seus prognósticos (resultado aos 90 minutos e equipa apurada) para cada jogo. "
     "O organizador regista as respostas na aplicação do Predictor.")
body(doc, "3.4  O envio dos prognósticos do mata-mata deve ser feito dentro do prazo "
     "comunicado pelo organizador para cada fase. Respostas fora do prazo equivalem a "
     "prognóstico em branco (0 pontos) para os jogos em causa.")
body(doc, "3.5  A validade dos prognósticos está condicionada à verificação descrita no Artigo 14.º.")
nota(doc, "Os prognósticos entregues em folha física têm precedência em caso de discrepância "
     "com a BD, desde que a situação seja reportada e devidamente documentada antes do apuramento "
     "final da classificação.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 4.º", "Fases do Torneio")
body(doc, "4.1  O Predictor divide-se em duas grandes fases, em espelho com o Campeonato do Mundo:")
bullet(doc, "Fase de Grupos — 48 jogos distribuídos por 12 grupos (A a L), com 3 jogos por grupo.")
bullet(doc, "Fase a Eliminar (Mata-Mata) — 24 jogos desde o Round of 32 (16 avos) até à Final, "
       "incluindo o jogo de atribuição do 3.º e 4.º lugar.")
body(doc, "4.2  Em ambas as fases, cada participante pontua de forma independente por cada jogo, "
     "conforme as regras dos Artigos 5.º e 6.º.", sb=4)
body(doc, "4.3  As pontuações de ambas as fases acumulam para a classificação geral final.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 5.º", "Pontuação — Fase de Grupos")
body(doc, "5.1  Na fase de grupos, cada participante previu o resultado exato (golos casa – golos fora) "
     "de cada jogo. A pontuação é atribuída por categorias não cumulativas:")
bullet(doc, "Resultado Exato (Exato):  5 pontos — Acertou o marcador exato do jogo.")
bullet(doc, "Vencedor ou Empate (VE):  2 pontos — Acertou quem ganhou ou que era empate, mas não o resultado exato.")
bullet(doc, "Golos de Uma Equipa:      1 ponto  — Acertou o número de golos de pelo menos uma das equipas, mas não o vencedor.")
bullet(doc, "Não Pontuou:              0 pontos — Não acertou nenhum dos critérios acima.")
body(doc, "5.2  As categorias não acumulam. O participante recebe apenas a pontuação da "
     "melhor categoria em que se enquadra.", sb=4, sa=6)

exemplo(doc, "Resultado Exato", [
    ("Real:", "França 3–1 Senegal"),
    ("Previsão:", "França 3–1 Senegal"),
    ("✅ Resultado exato →", "5 pontos"),
])
exemplo(doc, "Acertou o Vencedor", [
    ("Real:", "França 3–1 Senegal"),
    ("Previsão:", "França 2–0 Senegal"),
    ("✅ Vencedor correto · ❌ Não exato →", "2 pontos"),
])
exemplo(doc, "Acertou os Golos de Uma Equipa", [
    ("Real:", "França 3–1 Senegal"),
    ("Previsão:", "França 1–1 Senegal  (apostou empate)"),
    ("❌ Vencedor errado · ✅ Golos do Senegal: 1 →", "1 ponto"),
])
exemplo(doc, "Não Pontuou", [
    ("Real:", "França 3–1 Senegal"),
    ("Previsão:", "França 0–2 Senegal"),
    ("❌ Nada correto →", "0 pontos"),
])
exemplo(doc, "Pontos Não Acumulam — Regra Fundamental", [
    ("Real:", "França 3–1 Senegal"),
    ("Previsão:", "França 3–0 Senegal"),
    ("✅ Vencedor correto · ✅ Golos da França (3) corretos · ❌ Não exato", ""),
    ("→ Fica com a melhor categoria (VE):", "2 pontos   ⚠ NÃO soma 2+1=3"),
])

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 6.º", "Pontuação — Mata-Mata")
body(doc, "6.1  No mata-mata, além do resultado aos 90 minutos, cada participante seleciona "
     "a equipa que avança (apurado), o vencedor do 3.º lugar ou o campeão do mundo, "
     "consoante o jogo em causa.")
body(doc, "6.2  A pontuação do mata-mata divide-se em dois componentes independentes que acumulam:")
bullet(doc, "Score aos 90 minutos — avaliado pelas mesmas categorias da fase de grupos "
       "(Exato / VE / Golos / Nada), com pontos progressivamente mais altos por fase.")
bullet(doc, "Apurado / Campeão / 3.º Lugar — pontos atribuídos independentemente, "
       "se o participante indicou corretamente quem avança, independentemente do resultado.")
body(doc, "6.3  O resultado considerado é SEMPRE o resultado dos 90 minutos, mesmo que o jogo "
     "se prolongue para prolongamento ou penáltis. O prolongamento e os penáltis determinam "
     "apenas a equipa apurada.", sb=4)
nota(doc, "Se o jogo terminar 1–1 aos 90' e o participante previu 1–1 com Portugal a passar, "
     "e Portugal efetivamente passa (mesmo que nos penáltis), o participante recebe os pontos "
     "de resultado exato E os pontos de apurado — o máximo possível para essa fase.")
body(doc, "6.4  Os jogos do mata-mata só contam para pontuação após o participante ter registado "
     "os seus prognósticos na aplicação do Predictor. Prognósticos em falta equivalem a 0 pontos.", sb=4)

exemplo(doc, "Oitavos — Acertou Tudo", [
    ("Real:", "França 2–0 Croácia  |  França passa"),
    ("Previsão:", "França 2–0 Croácia  |  Apurado: França"),
    ("✅ Resultado exato: 10 pts  +  ✅ Apurado: 5 pts", ""),
    ("→ Total:", "15 pontos  (máximo nos Oitavos)"),
])
exemplo(doc, "Oitavos — Acertou Golos e Apurado", [
    ("Real:", "França 2–0 Croácia  |  França passa"),
    ("Previsão:", "França 2–2 Croácia  |  Apurado: França"),
    ("✅ Golos da França (2): 2 pts  +  ✅ Apurado: 5 pts", ""),
    ("→ Total:", "7 pontos"),
])
exemplo(doc, "Meias-Final — Empate e Penáltis", [
    ("Real:", "Brasil 1–1 França  (90')  |  França passa nos penáltis"),
    ("Previsão:", "Brasil 1–1 França  |  Apurado: França"),
    ("✅ Resultado exato (90'): 20 pts  +  ✅ Apurado: 15 pts", ""),
    ("→ Total:", "35 pontos  (máximo nas Meias-Finais)"),
])

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 7.º", "Tabela de Pontuação Progressiva")
body(doc, "7.1  A pontuação aumenta progressivamente à medida que o torneio avança, "
     "refletindo a crescente importância de cada jogo:")
tabela_pontuacao(doc)
body(doc, "7.2  O máximo teórico da competição é de 50 pontos na Final. "
     "O máximo teórico global é calculado somando o máximo de todos os jogos.", sb=4)
body(doc, "7.3  VE = Vencedor ou Empate. Golos = Acertou golos de uma equipa. "
     "Apurado = equipa apurada, campeão ou 3.º lugar (consoante o jogo).")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 8.º", "Classificação Geral")
body(doc, "8.1  A classificação geral é determinada pela soma total de pontos obtidos em todos "
     "os 72 jogos da fase de grupos e em todos os jogos do mata-mata em que o participante "
     "registou prognóstico.")
body(doc, "8.2  A classificação é apresentada na aplicação do Predictor em tempo real, "
     "sendo atualizada automaticamente após cada resultado introduzido.")
body(doc, "8.3  A classificação final é apurada após o resultado oficial da Final do Mundial 2026.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 9.º", "Critérios de Desempate")
body(doc, "9.1  Em caso de igualdade de pontos entre dois ou mais participantes, os critérios "
     "de desempate são aplicados pela seguinte ordem:")
numbered(doc, "1", "Maior número de resultados exatos (Exato) na fase de grupos.")
numbered(doc, "2", "Maior número de Vencedor/Empate na fase de grupos.")
numbered(doc, "3", "Maior pontuação obtida nos jogos do mata-mata.")
numbered(doc, "4", "Maior número de resultados exatos no mata-mata.")
numbered(doc, "5", "Maior pontuação obtida nos jogos do último grupo de jogos disputados.")
numbered(doc, "6", "Sorteio público, realizado na presença de todos os participantes empatados.")
body(doc, "9.2  Em nenhuma circunstância dois participantes partilham a mesma posição final. "
     "O desempate é sempre resolvido.", sb=4)
nota(doc, "O desempate por sorteio (critério 6) é o único critério que envolve aleatoriedade. "
     "Todos os outros critérios são objetivos e verificáveis na Base de Dados.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 10.º", "Prémio e Penalização — O Jantar")
body(doc, "10.1  O prémio do Predictor é o reconhecimento público dos 5 melhores classificados "
     "e, sobretudo, o jantar pago pelos 5 piores. A competição divide-se entre vencedores e pagadores.")
body(doc, "10.2  Sistema de pagamento do jantar:", sb=4)
bullet(doc, "Os 5 participantes classificados nas posições 6.ª a 10.ª (os últimos 5) pagam o jantar.")
bullet(doc, "Os 5 participantes classificados nas posições 1.ª a 5.ª (os primeiros 5) jantam gratuitamente.")
bullet(doc, "O custo total do jantar é dividido em partes iguais pelos 5 últimos classificados.")
body(doc, "10.3  O jantar inclui, por cada participante:", sb=4)
bullet(doc, "Entrada.")
bullet(doc, "Prato principal.")
bullet(doc, "Bebidas (água, refrigerantes, vinho ou cerveja — incluídas na refeição).")
bullet(doc, "Sobremesa.")
body(doc, "10.4  Bebidas adicionais ou fora do menu combinado são pagas individualmente, "
     "pelos vencedores ou pelos pagadores, consoante quem as consuma.", sb=4)
body(doc, "10.5  O restaurante é escolhido por votação simples (maioria) dos 5 primeiros classificados. "
     "Os 5 últimos classificados não têm direito de voto na escolha do local.")
body(doc, "10.6  O 1.º classificado tem direito de escolher a data do jantar, "
     "em coordenação com a disponibilidade geral do grupo.", sb=4)
body(doc, "10.7  Cada um dos 5 últimos classificados paga a sua quota-parte no próprio dia do jantar, "
     "antes de sair do restaurante. Não são aceites transferências diferidas salvo acordo "
     "expresso dos 5 primeiros classificados.")

exemplo(doc, "Cenário do Jantar", [
    ("1.º ao 5.º lugar:", "André, João Garcia, Hélio, Miguel, Luis — jantam de graça."),
    ("6.º ao 10.º lugar:", "Vítor, Paulo, Ricardo, Rúben, José — pagam o jantar."),
    ("Divisão do custo:", "Conta total ÷ 5 = quota de cada um dos últimos 5."),
    ("Escolha do restaurante:", "Votação dos 5 primeiros (os últimos não votam)."),
    ("Data:", "Decidida pelo 1.º classificado, com consulta geral."),
])

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 11.º", "Obrigações dos Últimos 5 Classificados")
body(doc, "11.1  Cada um dos 5 últimos classificados (posições 6.ª a 10.ª) assume as seguintes obrigações:")
bullet(doc, "Confirmar a disponibilidade para o jantar nas primeiras 48 horas após a Final do Mundial.")
bullet(doc, "Aceitar o restaurante escolhido pelos 5 primeiros sem objeções.")
bullet(doc, "Pagar a sua quota-parte da conta no próprio dia do jantar.")
bullet(doc, "Não impor restrições orçamentais ao grupo na escolha do restaurante — dentro do razoável.")
bullet(doc, "Comparecer ao jantar (ver Artigo 12.º sobre ausência).")
body(doc, "11.2  Os últimos 5 classificados têm pleno direito de participar no jantar e desfrutar da "
     "noite como qualquer outro convidado. A penalização é financeira, não social.", sb=4)
body(doc, "11.3  Em caso de desistência de algum dos 5 primeiros (que jantam de graça), "
     "a sua vaga não é transferível nem gera redução no valor pago pelos últimos 5.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 12.º", "Ausência ao Jantar")
body(doc, "12.1  A presença no jantar é voluntária para os 5 primeiros classificados. "
     "Se algum dos 5 primeiros não puder comparecer, a sua quota é simplesmente não paga — "
     "os últimos 5 não ficam beneficiados por isso.")
body(doc, "12.2  Para os 5 ÚLTIMOS CLASSIFICADOS, a presença é obrigatória sob pena das "
     "sanções previstas neste artigo.", sb=4)
body(doc, "12.3  Se algum dos 5 últimos não comparecer SEM justificação válida "
     "(ver n.º 12.5), aplica-se:", sb=4)
bullet(doc, "Pagamento por transferência bancária da sua quota-parte ao 1.º classificado, "
       "que redistribui ao grupo consoante os custos reais.")
bullet(doc, "Publicação de mensagem de rendição no grupo de WhatsApp, "
       "reconhecendo publicamente a derrota.")
bullet(doc, "Obrigação de organizar um segundo jantar em data futura, a definir pelo grupo.")
body(doc, "12.4  Se vários dos 5 últimos faltarem sem justificação, cada um enfrenta "
     "as penalizações acima de forma individual e cumulativa.", sb=4)
body(doc, "12.5  Constituem justificações válidas para ausência:", sb=4)
bullet(doc, "Doença grave devidamente documentada (declaração médica).")
bullet(doc, "Emergência familiar de primeiro grau.")
bullet(doc, "Ausência do país por motivos profissionais inadiáveis.")
bullet(doc, "Qualquer outro motivo aceite por unanimidade dos 5 primeiros classificados.")
body(doc, "12.6  Em caso de ausência justificada de qualquer dos 5 últimos, "
     "o jantar é remarcado para data conveniente para a maioria. "
     "A obrigação de pagamento mantém-se.", sb=4)
nota(doc, "«Não me apetece sair» e «tenho muito trabalho» não constituem justificações válidas. "
     "O jantar é uma obrigação assumida no momento da inscrição e aceite neste regulamento.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 13.º", "Suspense — Revelação dos Resultados no Jantar")
body(doc, "13.1  Um dos momentos centrais do Predictor é a revelação pública e dramática da "
     "classificação final e dos prognósticos individuais de cada participante. "
     "Para preservar o suspense e maximizar o entretenimento no jantar, aplicam-se as "
     "seguintes regras de confidencialidade.")
body(doc, "13.2  Até ao momento da revelação no jantar:", sb=4)
bullet(doc, "Nenhum participante conhece os prognósticos detalhados dos outros participantes.")
bullet(doc, "O organizador não divulga a classificação provisória nem os pontos individuais "
       "de cada participante a terceiros.")
bullet(doc, "A aplicação do Predictor pode ser consultada pelo organizador, "
       "mas os dados não são partilhados publicamente antes do jantar.")
body(doc, "13.3  A revelação dos resultados é feita a meio do jantar, num momento cerimonial "
     "organizado pelo 1.º classificado ou pelo organizador, da seguinte forma sugerida:", sb=4)
numbered(doc, "1", "Anúncio da pontuação total de cada participante, do 10.º para o 1.º lugar.")
numbered(doc, "2", "Revelação dos prognósticos mais marcantes (mais exatos, mais errados, mais ousados).")
numbered(doc, "3", "Entrega simbólica do troféu ou reconhecimento ao 1.º classificado.")
numbered(doc, "4", "Momento de 'shame' saudável dos 5 últimos classificados, que nessa altura já sabem quem paga.")
body(doc, "13.4  A ordem e o formato da revelação podem ser adaptados pelo organizador "
     "consoante o ambiente e dinâmica do grupo no jantar.", sb=4)
body(doc, "13.5  É expressamente proibido divulgar a classificação final ou os prognósticos "
     "detalhados de outros participantes antes do jantar, seja por WhatsApp, redes sociais "
     "ou qualquer outro meio. Quem o fizer estraga o suspense e merece olhares de reprovação geral.")
nota(doc, "O suspense é parte integrante da experiência. Saber que pode estar nos últimos 5 "
     "mas não saber ao certo até ao jantar torna a noite muito mais especial. "
     "Respeitem o processo.")

exemplo(doc, "Momento da Revelação no Jantar", [
    ("A meio do jantar:", "O organizador abre o portátil/telemóvel com a aplicação do Predictor."),
    ("Anúncio dramático:", "«Em 10.º lugar, com X pontos... [pausa]... FULANO!»"),
    ("Reação esperada:", "Gargalhadas, espanto, reclamações amigáveis, alívio dos vencedores."),
    ("Clímax:", "Revelação do 1.º classificado e dos prognósticos mais certeiros."),
])

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 14.º", "Reclamações e Resolução de Disputas")
body(doc, "13.1  Qualquer participante que discorde de uma pontuação atribuída, de um resultado "
     "introduzido, ou de qualquer decisão operacional do Predictor, pode apresentar reclamação "
     "formal nos termos deste artigo.")
body(doc, "13.2  Prazo de reclamação:", sb=4)
bullet(doc, "Fase de Grupos: a reclamação deve ser apresentada até 24 horas após "
       "a introdução do resultado em disputa na aplicação.")
bullet(doc, "Mata-Mata: a reclamação deve ser apresentada até 12 horas após "
       "a introdução do resultado em disputa.")
bullet(doc, "Após estes prazos, o resultado considera-se aceite e definitivo.")
body(doc, "13.3  A reclamação deve:", sb=4)
bullet(doc, "Ser comunicada diretamente ao organizador (João Garcia) por mensagem escrita "
       "(WhatsApp ou equivalente).")
bullet(doc, "Identificar claramente o jogo em disputa (código do jogo), o resultado introduzido "
       "e o resultado que o participante considera correto.")
bullet(doc, "Incluir evidência de suporte (ex.: captura de ecrã do resultado oficial FIFA, "
       "link para a notícia, ou referência à folha original de prognósticos).")
body(doc, "13.4  Processo de resolução:", sb=4)
numbered(doc, "1", "O organizador analisa a reclamação e verifica com fontes oficiais (FIFA, UEFA, etc.).")
numbered(doc, "2", "A decisão é comunicada a todos os participantes no prazo de 24 horas.")
numbered(doc, "3", "Se a reclamação proceder, o resultado é corrigido e a classificação atualizada.")
numbered(doc, "4", "Se a reclamação não proceder, o organizador explica os motivos por escrito.")
numbered(doc, "5", "Não existe recurso da decisão do organizador. A sua palavra é final.")
body(doc, "13.5  Reclamações manifestamente infundadas, feitas de má-fé ou com caráter "
     "perturbador repetido, podem resultar em aviso formal. Após dois avisos, o organizador "
     "pode ignorar reclamações futuras do mesmo participante.", sb=4)
nota(doc, "Reclamações sobre a sorte («eu devia ter acertado aquele jogo»), sobre as prestações "
     "das seleções («não era suposto o Brasil perder»), ou sobre regras claramente estabelecidas "
     "neste regulamento não serão consideradas. O regulamento é claro e foi aceite na inscrição.")

exemplo(doc, "Reclamação Válida", [
    ("Situação:", "O organizador introduziu Portugal 2–1 França, mas o resultado real foi 2–2."),
    ("Ação:", "Participante envia captura do resultado oficial FIFA + link da notícia."),
    ("Resolução:", "Organizador verifica, confirma o erro, corrige o resultado e a classificação."),
    ("Prazo:", "Resolução em até 24 horas após a reclamação."),
])
exemplo(doc, "Reclamação Inválida", [
    ("Situação:", "Participante diz: «eu previa que o Brasil ia ganhar, não faz sentido.»"),
    ("Ação do organizador:", "Reclamação rejeitada por não ter fundamento objetivo."),
    ("Nota:", "Sentimentos e opiniões não alteram resultados."),
])

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 15.º", "Base de Dados e Verificação dos Prognósticos")
body(doc, "14.1  A Base de Dados (BD) digital do Predictor é a fonte oficial de todos os "
     "prognósticos e resultados. É gerida pelo organizador e acessível a todos os participantes "
     "através da aplicação local.")
body(doc, "14.2  Antes do apuramento da classificação final, cada participante recebe uma "
     "Folha de Verificação de Prognósticos — um documento Word com todos os seus 72 prognósticos "
     "da fase de grupos registados na BD — para confrontar com a fotocópia da sua folha original.")
body(doc, "14.3  Se um participante detetar discrepâncias entre a sua folha original e a BD, "
     "deve comunicá-las ao organizador até ao prazo de reclamação aplicável (Artigo 13.º).", sb=4)
body(doc, "14.4  A folha original entregue na inscrição tem precedência sobre a BD em caso "
     "de discrepância comprovada. O participante deve apresentar a folha original ou cópia "
     "legível como evidência.")
body(doc, "15.5  Os prognósticos do mata-mata são enviados digitalmente pelo participante "
     "ao organizador (conforme Artigo 3.º) e registados na aplicação. "
     "O registo é definitivo e irrevogável após o início do jogo a que se referem.", sb=4)
nota(doc, "A BD é apenas tão boa quanto os dados introduzidos. Qualquer discrepância deve ser "
     "reportada imediatamente. Após o prazo de reclamação, os dados da BD são definitivos.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 16.º", "Alteração de Prognósticos")
body(doc, "15.1  Os prognósticos da fase de grupos são registados antes do início do torneio "
     "e são imutáveis após o arranque da competição, salvo no âmbito do processo de verificação "
     "descrito no Artigo 14.º.")
body(doc, "15.2  Os prognósticos do mata-mata podem ser registados ou alterados a qualquer "
     "momento até ao apito inicial do jogo em questão.")
body(doc, "15.3  Após o apito inicial de um jogo, o prognóstico para esse jogo específico "
     "fica bloqueado e não pode ser alterado por qualquer razão.", sb=4)
body(doc, "15.4  Alterações feitas na aplicação após o início do jogo não são reconhecidas "
     "e não afetam a pontuação, mesmo que o organizador as detete e as corrija retroativamente.")
nota(doc, "«Não consegui registar a tempo» não é justificação para alterar um prognóstico "
     "após o início do jogo. Cada participante é responsável por registar os seus prognósticos "
     "atempadamente. A gestão do tempo é parte integrante da competição.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 17.º", "Jogos Cancelados, Adiados ou com Resultado Anulado")
body(doc, "16.1  Se um jogo for adiado mas realizado dentro do calendário oficial da FIFA, "
     "o resultado é introduzido na data efetiva do jogo. Os prognósticos mantêm-se válidos.")
body(doc, "16.2  Se um jogo for cancelado definitivamente pela FIFA ou equivalente, "
     "esse jogo é retirado do Predictor e não conta para a pontuação de nenhum participante. "
     "A classificação é recalculada sem esse jogo.", sb=4)
body(doc, "16.3  Se um resultado for anulado por decisão disciplinar (ex.: resultado de secretaria "
     "por infração de elegibilidade de um jogador), é o resultado oficial pós-anulação "
     "(normalmente 3–0) que é considerado para o Predictor.")
body(doc, "16.4  Resultados alterados por decisão de VAR ou árbitro durante o jogo fazem parte "
     "do jogo normal e são considerados parte do resultado final. Não há tratamento especial.", sb=4)
nota(doc, "O Predictor segue o resultado oficial final reconhecido pela FIFA, "
     "independentemente das circunstâncias que levaram a esse resultado.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 18.º", "Conduta e Fair-Play")
body(doc, "17.1  Todos os participantes comprometem-se a participar com espírito de fair-play, "
     "honestidade e boa disposição.")
body(doc, "17.2  São consideradas infrações graves:")
bullet(doc, "Tentativa de aceder ou alterar a Base de Dados sem autorização do organizador.")
bullet(doc, "Falsificação ou alteração de folhas originais de prognósticos.")
bullet(doc, "Coação, ameaça ou pressão sobre o organizador para alterar resultados ou classificações.")
bullet(doc, "Divulgação de informação privilegiada sobre prognósticos alheios sem consentimento.")
body(doc, "17.3  Em caso de infração grave comprovada, o organizador pode aplicar, "
     "por ordem crescente de gravidade:", sb=4)
bullet(doc, "Aviso formal por escrito.")
bullet(doc, "Anulação dos pontos do(s) jogo(s) em causa.")
bullet(doc, "Desqualificação do participante do Predictor, com perda de todos os pontos.")
body(doc, "17.4  A desqualificação não isenta o participante da obrigação do jantar, "
     "caso tenha terminado em último lugar antes da desqualificação.", sb=4)

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 19.º", "Privacidade e Uso dos Dados")
body(doc, "18.1  Os dados pessoais dos participantes (nome e prognósticos) são utilizados "
     "exclusivamente no âmbito do Predictor, para efeitos de pontuação, classificação e "
     "geração de documentos de verificação.")
body(doc, "18.2  A aplicação do Predictor funciona exclusivamente em modo local (localhost), "
     "sem ligação a servidores externos, sem armazenamento em cloud e sem transmissão "
     "de dados a terceiros.", sb=4)
body(doc, "18.3  Os dados são armazenados no dispositivo do organizador através do mecanismo "
     "localStorage do browser. O organizador assume responsabilidade pela guarda e "
     "confidencialidade dos dados.")
body(doc, "18.4  Os prognósticos de cada participante não são divulgados a outros participantes "
     "durante a competição, para preservar a integridade competitiva. "
     "A divulgação pode ser feita no final do torneio, com consentimento de todos.")

# ──────────────────────────────────────────────────────────────────────────────
art_title(doc, "Artigo 20.º", "Disposições Finais")
body(doc, "19.1  O presente regulamento entra em vigor no momento da inscrição do participante "
     "e mantém-se válido até ao encerramento completo do Predictor, incluindo a realização do jantar.")
body(doc, "19.2  O organizador reserva-se o direito de clarificar, interpretar e, em casos "
     "excecionais não previstos neste regulamento, tomar decisões ad hoc, sempre em espírito "
     "de equidade e bom senso.", sb=4)
body(doc, "19.3  Qualquer alteração a este regulamento após o início da competição requer "
     "a concordância de, pelo menos, 7 dos 10 participantes.")
body(doc, "19.4  O espírito do Predictor é o de convívio, amizade e saudável competição. "
     "As regras existem para garantir equidade, não para criar conflito. Em caso de dúvida, "
     "prevalecem sempre os valores da amizade e do bom senso.", sb=4)
body(doc, "19.5  O não cumprimento do regulamento não gera qualquer obrigação jurídica entre "
     "as partes. A única sanção prevista é de caráter social e simbólico (o jantar). "
     "Este documento não tem qualquer valor legal.")

# ──────────────────────────────────────────────────────────────────────────────
# ASSINATURA
# ──────────────────────────────────────────────────────────────────────────────
br3 = doc.add_paragraph(); fmt(br3)
br3.add_run().add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)

for _ in range(3): doc.add_paragraph()

sign_title = doc.add_paragraph()
fmt(sign_title, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
run(sign_title, "ASSINATURA DO ORGANIZADOR", bold=True, sz=11, color=AZUL)

sign_sep = doc.add_paragraph()
fmt(sign_sep, align=WD_ALIGN_PARAGRAPH.CENTER, sa=8)
run(sign_sep, "─" * 40, sz=6, color=AZUL)

sign_body = doc.add_paragraph()
fmt(sign_body, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
run(sign_body, "O presente Regulamento foi elaborado, aprovado e é da inteira responsabilidade de:", sz=10, color=CZ2)

for _ in range(2): doc.add_paragraph()

sign_line = doc.add_paragraph()
fmt(sign_line, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(sign_line, "_" * 50, sz=11)

sign_name = doc.add_paragraph()
fmt(sign_name, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(sign_name, "João Garcia", bold=True, sz=15, color=AZUL)

sign_role = doc.add_paragraph()
fmt(sign_role, align=WD_ALIGN_PARAGRAPH.CENTER, sa=1)
run(sign_role, "Organizador · Predictor Parque Biológico — Mundial 2026", italic=True, sz=10, color=CZ2)

sign_date = doc.add_paragraph()
fmt(sign_date, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
run(sign_date, f"Lisboa, {hoje}", sz=10, color=CZ2)

for _ in range(4): doc.add_paragraph()

final_note = doc.add_paragraph()
fmt(final_note, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
run(final_note, "«O futebol é imprevisível. Os prognósticos, ainda mais.", italic=True, sz=10, color=AZUL)
run(final_note, "\nMas as regras do Predictor — essas são definitivas.»", italic=True, sz=10, color=AZUL)

for _ in range(2): doc.add_paragraph()

ver_p = doc.add_paragraph()
fmt(ver_p, align=WD_ALIGN_PARAGRAPH.CENTER, sa=0)
run(ver_p, f"Versão 1.0  ·  Emitido em {hoje}  ·  Predictor Parque Biológico — Mundial 2026",
    sz=8, color=RGBColor(0xbb,0xbb,0xbb))

# ── Guardar ───────────────────────────────────────────────────────────────────
OUT = "/Users/joaogarcia/Projects/predictor-mundial-2026/Regulamento_Predictor_Mundial_2026.docx"
doc.save(OUT)
print(f"✅  {OUT}")
